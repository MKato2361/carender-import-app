# tabs/tab8_notice_fax.py
from __future__ import annotations

from datetime import datetime, date, time, timedelta, timezone
from typing import Any, Dict, Optional, List

import io
import zipfile
import unicodedata
import re
from pathlib import Path

import pandas as pd
import streamlit as st
from firebase_admin import firestore
from docx import Document

from github_loader import load_file_bytes_from_github
from tabs.tab6_property_master import (
    MASTER_COLUMNS,
    BASIC_COLUMNS,
    load_sheet_as_df,
    _normalize_df,
)
from utils.helpers import safe_get

# ==========================
# 定数
# ==========================
JST = timezone(timedelta(hours=9))
WEEKDAYS_JA = "月火水木金土日"

# デフォルトテンプレート（リポジトリ内パス想定）
DEFAULT_NOTICE_TEMPLATE_PATH = "templates/harigami.docx"
DEFAULT_FAX_TEMPLATE_PATH = "templates/fax_template.docx"

# 管理番号抽出
ASSETNUM_PATTERN = re.compile(
    r"[［\[]?\s*管理番号[：:]\s*([0-9A-Za-z\-]+)\s*[］\]]?"
)


# ==========================
# 共通小物
# ==========================
def normalize_str(val: Any) -> str:
    if val is None:
        return ""
    return unicodedata.normalize("NFKC", str(val)).strip()


def display_value(val: Any) -> str:
    """nan / 空 → '-' にそろえる"""
    if val is None:
        return "-"
    s = str(val).strip()
    if not s:
        return "-"
    if s.lower() in ("nan", "none"):
        return "-"
    return s


def format_date_ja(d: date) -> str:
    """例: 2025-01-23 → '1月23日（木）'"""
    w = WEEKDAYS_JA[d.weekday()]
    return f"{d.month}月{d.day}日（{w}）"


def extract_assetnum(text: str) -> str:
    """Description 等から管理番号を抽出"""
    if not text:
        return ""
    s = normalize_str(text)
    m = ASSETNUM_PATTERN.search(s)
    if not m:
        return ""
    return m.group(1).strip()


def get_event_start_datetime(event: Dict[str, Any]) -> Optional[datetime]:
    """Googleカレンダーイベントから開始日時（JST）を取得"""
    start = event.get("start", {})
    # 時間付き予定
    if "dateTime" in start:
        try:
            dt = pd.to_datetime(start["dateTime"])
            if dt.tzinfo is None:
                dt = dt.tz_localize(timezone.utc)
            dt = dt.astimezone(JST)
            return dt.to_pydatetime()
        except Exception:
            return None
    # 終日予定
    if "date" in start:
        try:
            d = date.fromisoformat(start["date"])
            return datetime.combine(d, time.min, tzinfo=JST)
        except Exception:
            return None
    return None


def get_event_end_datetime(event: Dict[str, Any]) -> Optional[datetime]:
    """終了日時（JST）を取得（なければ None）"""
    end = event.get("end", {})
    # 時間付き予定
    if "dateTime" in end:
        try:
            dt = pd.to_datetime(end["dateTime"])
            if dt.tzinfo is None:
                dt = dt.tz_localize(timezone.utc)
            dt = dt.astimezone(JST)
            return dt.to_pydatetime()
        except Exception:
            return None
    # 終日予定はここでは特に使わないので None のままでもOK
    return None


def to_utc_range_from_dates(d1: date, d2: date) -> tuple[str, str]:
    """date 範囲 → Calendar API 用 UTC 時刻範囲"""
    start_dt_utc = datetime.combine(d1, time.min, tzinfo=JST).astimezone(timezone.utc)
    end_dt_utc = datetime.combine(d2, time.max, tzinfo=JST).astimezone(timezone.utc)
    return start_dt_utc.isoformat(), end_dt_utc.isoformat()


def fetch_events_in_range(
    service: Any,
    calendar_id: str,
    start_date: date,
    end_date: date,
) -> List[Dict[str, Any]]:
    """指定カレンダーのイベントを期間指定で全取得"""
    if not service:
        return []

    time_min, time_max = to_utc_range_from_dates(start_date, end_date)

    events: List[Dict[str, Any]] = []
    page_token: Optional[str] = None

    while True:
        resp = (
            service.events()
            .list(
                calendarId=calendar_id,
                timeMin=time_min,
                timeMax=time_max,
                maxResults=2500,
                singleEvents=True,
                orderBy="startTime",
                pageToken=page_token,
            )
            .execute()
        )
        items = resp.get("items", [])
        events.extend(items)
        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    return events


# ==========================
# 物件マスタ関連（tab7 と同じロジック）
# ==========================
def get_property_master_spreadsheet_id(current_user_email: Optional[str]) -> str:
    """Firestore の user_settings からユーザーごとの物件マスタ用 Spreadsheet ID を取得"""
    if not current_user_email:
        return ""
    try:
        db = firestore.client()
        doc = db.collection("user_settings").document(current_user_email).get()
        if not doc.exists:
            return ""
        data = doc.to_dict() or {}
        return data.get("property_master_spreadsheet_id") or ""
    except Exception as e:
        st.warning(f"物件マスタ用スプレッドシートIDの取得に失敗しました: {e}")
        return ""


def load_property_master_view(
    sheets_service: Any,
    spreadsheet_id: str,
    basic_sheet_title: str = "物件基本情報",
    master_sheet_title: str = "物件マスタ",
) -> pd.DataFrame:
    """
    物件基本情報（BASIC_COLUMNS）＋物件マスタ（MASTER_COLUMNS）を読み込み、
    管理番号でマージした DataFrame を返す。
    """
    if not sheets_service or not spreadsheet_id:
        return pd.DataFrame()

    try:
        basic_df = load_sheet_as_df(
            sheets_service,
            spreadsheet_id,
            basic_sheet_title,
            BASIC_COLUMNS,
        )
        master_df = load_sheet_as_df(
            sheets_service,
            spreadsheet_id,
            master_sheet_title,
            MASTER_COLUMNS,
        )
    except Exception as e:
        st.error(f"物件マスタの読み込みに失敗しました: {e}")
        return pd.DataFrame()

    basic_df = _normalize_df(basic_df, BASIC_COLUMNS)
    master_df = _normalize_df(master_df, MASTER_COLUMNS)

    if master_df.empty:
        # マスタが空なら管理番号だけでも返しておく
        merged = basic_df.copy()
        for col in MASTER_COLUMNS:
            if col not in merged.columns:
                merged[col] = ""
        return merged

    merged = master_df.merge(
        basic_df[["管理番号", "物件名", "住所", "窓口会社"]],
        on="管理番号",
        how="left",
    )

    return merged


# ==========================
# テンプレート関連
# ==========================
def load_template_bytes(default_path: str, uploaded_file) -> Optional[bytes]:
    """
    DOCX テンプレートのバイト列を取得。
    - アップロードがあればそれを優先
    - なければ GitHub 上の default_path を load_file_bytes_from_github で取得
    - それも失敗したらローカルファイル default_path を読む
    """
    if uploaded_file is not None:
        return uploaded_file.getvalue()

    # GitHub から
    try:
        return load_file_bytes_from_github(default_path)
    except Exception:
        pass

    # ローカル fallback
    p = Path(default_path)
    if p.exists():
        return p.read_bytes()

    return None


def replace_placeholders_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """段落内のプレースホルダを単純置換"""
    if not replacements:
        return
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)


def replace_placeholders_in_doc(doc: Document, replacements: Dict[str, str]) -> None:
    """文書全体（本文＋表）のプレースホルダを置換"""
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, replacements)


def build_notice_replacements(
    ev: Dict[str, Any],
    pm_row: pd.Series,
    start_dt: Optional[datetime],
    end_dt: Optional[datetime],
    company_name: str,
    company_tel: str,
) -> Dict[str, str]:
    """貼り紙テンプレ用のプレースホルダ置換辞書"""
    mgmt = display_value(pm_row.get("管理番号", ""))
    name = display_value(pm_row.get("物件名", ""))
    addr = display_value(pm_row.get("住所", ""))

    if start_dt:
        date_label = format_date_ja(start_dt.date())
        date_full = start_dt.strftime("%Y年%m月%d日")
    else:
        date_label = "-"
        date_full = "-"

    # 時間帯
    start_time_str = ""
    end_time_str = ""
    time_range_str = "-"
    if start_dt and "dateTime" in ev.get("start", {}):
        start_time_str = start_dt.strftime("%H:%M")
        if end_dt and "dateTime" in ev.get("end", {}):
            end_time_str = end_dt.strftime("%H:%M")
            time_range_str = f"{start_time_str}〜{end_time_str}"
        else:
            time_range_str = start_time_str
    elif start_dt and "date" in ev.get("start", {}):
        time_range_str = "終日"

    reps = {
        "［物件名］": name,
        "［管理番号］": mgmt,
        "［住所］": addr,
        "［点検予定］": f"{date_label} {time_range_str}" if date_label != "-" else "-",
        "［点検日］": date_full,
        "［点検時間］": time_range_str,
        "［会社名］": company_name,
        "［自社名］": company_name,
        "［会社TEL］": company_tel,
        "［自社TEL］": company_tel,
    }

    # イベントタイトルを使いたい場合用
    event_title = display_value(safe_get(ev, "summary") or "")
    reps["［イベントタイトル］"] = event_title

    return reps


def build_fax_replacements(
    ev: Dict[str, Any],
    pm_row: pd.Series,
    start_dt: Optional[datetime],
    end_dt: Optional[datetime],
    company_name: str,
    company_tel: str,
    company_fax: str,
    contact_person: str,
) -> Dict[str, str]:
    """FAXテンプレ用のプレースホルダ置換辞書"""
    mgmt = display_value(pm_row.get("管理番号", ""))
    name = display_value(pm_row.get("物件名", ""))
    addr = display_value(pm_row.get("住所", ""))

    dest1 = display_value(pm_row.get("連絡宛名1", ""))
    dest2 = display_value(pm_row.get("連絡宛名2", ""))
    fax1 = display_value(pm_row.get("FAX番号1", ""))
    fax2 = display_value(pm_row.get("FAX番号2", ""))

    if start_dt:
        date_label = format_date_ja(start_dt.date())
        date_full = start_dt.strftime("%Y年%m月%d日")
    else:
        date_label = "-"
        date_full = "-"

    # 時間帯
    start_time_str = ""
    end_time_str = ""
    time_range_str = "-"
    if start_dt and "dateTime" in ev.get("start", {}):
        start_time_str = start_dt.strftime("%H:%M")
        if end_dt and "dateTime" in ev.get("end", {}):
            end_time_str = end_dt.strftime("%H:%M")
            time_range_str = f"{start_time_str}〜{end_time_str}"
        else:
            time_range_str = start_time_str
    elif start_dt and "date" in ev.get("start", {}):
        time_range_str = "終日"

    today_str = datetime.now(JST).strftime("%Y年%m月%d日")

    event_title = display_value(safe_get(ev, "summary") or "")

    reps = {
        "［宛名1］": dest1,
        "［宛名２］": dest2,
        "［宛名2］": dest2,
        "［FAX番号1］": fax1,
        "［FAX番号２］": fax2,
        "［FAX番号2］": fax2,
        "［物件名］": name,
        "［管理番号］": mgmt,
        "［住所］": addr,
        "［点検予定］": f"{date_label} {time_range_str}" if date_label != "-" else "-",
        "［点検日］": date_full,
        "［点検時間］": time_range_str,
        "［送信日］": today_str,
        "［会社名］": company_name,
        "［自社名］": company_name,
        "［会社TEL］": company_tel,
        "［自社TEL］": company_tel,
        "［会社FAX］": company_fax,
        "［自社FAX］": company_fax,
        "［担当者名］": contact_person,
        "［イベントタイトル］": event_title,
    }

    return reps


def build_zip_from_files(files: List[tuple[str, bytes]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname, content in files:
            # 同名ファイルがあったら少しずらす
            safe_name = fname
            i = 1
            while safe_name in zf.namelist():
                stem = Path(fname).stem
                suffix = Path(fname).suffix
                safe_name = f"{stem}_{i}{suffix}"
                i += 1
            zf.writestr(safe_name, content)
    buf.seek(0)
    return buf.getvalue()


# ==========================
# メイン UI
# ==========================
def render_tab8_notice_fax(
    service: Any,
    editable_calendar_options: Dict[str, str],
    sheets_service: Any,
    current_user_email: Optional[str] = None,
) -> None:
    """
    貼り紙・FAX 自動生成タブ
      - 指定期間のイベント取得
      - 物件マスタと突合
      - 貼り紙テンプレ種別 = '自社' → 貼り紙 DOCX 生成
      - 連絡方法_FAX1/2 が設定されている → FAX DOCX 生成
      - 生成した DOCX を ZIP で一括ダウンロード
    """
    st.subheader("貼り紙・FAX 自動生成")

    if not service or not editable_calendar_options:
        st.warning("カレンダーサービスが初期化されていません。先にGoogle認証を完了してください。")
        return

    # Sheets サービス必須
    if not sheets_service:
        st.info(
            "この機能を利用するには、まずタブ『物件マスタ管理』で\n"
            "物件マスタ用スプレッドシートを作成してください。"
        )
        return

    # 物件マスタ用スプレッドシートID
    spreadsheet_id = get_property_master_spreadsheet_id(current_user_email)
    if not spreadsheet_id:
        st.info(
            "物件マスタがまだ設定されていません。\n\n"
            "タブ『物件マスタ管理』で物件マスタを作成し、保存すると\n"
            "貼り紙・FAX自動生成機能を利用できるようになります。"
        )
        return

    st.markdown(
        f"物件マスタスプレッドシート: "
        f"[リンクを開く](https://docs.google.com/spreadsheets/d/{spreadsheet_id})"
    )

    # 物件マスタビュー読み込み
    pm_view_df = load_property_master_view(
        sheets_service,
        spreadsheet_id,
        basic_sheet_title="物件基本情報",
        master_sheet_title="物件マスタ",
    )
    if pm_view_df is None or pm_view_df.empty:
        st.warning(
            "物件マスタ（＋基本情報）が空です。\n\n"
            "タブ『物件マスタ管理』で物件情報を登録すると、\n"
            "貼り紙・FAX自動生成機能を利用できるようになります。"
        )
        return

    st.caption(f"物件マスタ登録件数: {len(pm_view_df)} 件（管理番号単位）")

    # カレンダー選択
    cal_names = list(editable_calendar_options.keys())
    default_cal = cal_names[0] if cal_names else None
    calendar_name = st.selectbox(
        "対象カレンダー",
        cal_names,
        index=(cal_names.index(default_cal) if default_cal in cal_names else 0),
        key="notice_fax_calendar",
    )
    calendar_id = editable_calendar_options.get(calendar_name)

    # 期間指定
    today = date.today()
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        start_date = st.date_input("点検予定の検索開始日", value=today, key="notice_fax_start_date")
    with col_d2:
        end_date = st.date_input("点検予定の検索終了日", value=today + timedelta(days=60), key="notice_fax_end_date")

    if start_date > end_date:
        st.error("開始日は終了日以前の日付を指定してください。")
        return

    st.markdown("#### 生成対象とテンプレート")

    col_flag1, col_flag2 = st.columns(2)
    with col_flag1:
        generate_notice = st.checkbox("貼り紙を生成する", value=True, key="notice_generate")
    with col_flag2:
        generate_fax = st.checkbox("FAXを生成する", value=True, key="fax_generate")

    if not (generate_notice or generate_fax):
        st.info("貼り紙またはFAXのどちらかを生成対象として選択してください。")
        return

    st.markdown("##### 共通情報（テンプレートに差し込む自社情報など）")
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        company_name = st.text_input("自社名 / 会社名", value="〇〇エレベーター株式会社", key="nf_company_name")
        company_tel = st.text_input("自社電話番号（例: TEL: 03-xxxx-xxxx）", value="", key="nf_company_tel")
    with col_c2:
        company_fax = st.text_input("自社FAX番号（例: FAX: 03-xxxx-xxxx）", value="", key="nf_company_fax")
        contact_person = st.text_input("担当者名（FAX用）", value="", key="nf_contact_person")

    st.markdown("##### テンプレート選択")

    col_tpl1, col_tpl2 = st.columns(2)
    with col_tpl1:
        if generate_notice:
            notice_mode = st.radio(
                "貼り紙テンプレート",
                ["デフォルトを使用", "DOCXファイルをアップロード"],
                key="notice_tpl_mode",
            )
            notice_upload = None
            if notice_mode == "DOCXファイルをアップロード":
                notice_upload = st.file_uploader(
                    "貼り紙テンプレート（.docx）",
                    type=["docx"],
                    key="notice_tpl_upload",
                )
        else:
            notice_mode = "デフォルトを使用"
            notice_upload = None

    with col_tpl2:
        if generate_fax:
            fax_mode = st.radio(
                "FAXテンプレート",
                ["デフォルトを使用", "DOCXファイルをアップロード"],
                key="fax_tpl_mode",
            )
            fax_upload = None
            if fax_mode == "DOCXファイルをアップロード":
                fax_upload = st.file_uploader(
                    "FAXテンプレート（.docx）",
                    type=["docx"],
                    key="fax_tpl_upload",
                )
        else:
            fax_mode = "デフォルトを使用"
            fax_upload = None

    # 生成ボタン
    generate_btn = st.button("指定期間のイベントから貼り紙・FAX文書を作成する", type="primary", key="notice_fax_generate")

    if generate_btn:
        with st.spinner("カレンダーイベントの取得と文書生成を行っています..."):
            events = fetch_events_in_range(service, calendar_id, start_date, end_date)
            st.write(f"取得したイベント件数: {len(events)} 件")

            # テンプレート読み込み
            notice_tpl_bytes: Optional[bytes] = None
            fax_tpl_bytes: Optional[bytes] = None

            if generate_notice:
                notice_tpl_bytes = load_template_bytes(DEFAULT_NOTICE_TEMPLATE_PATH, notice_upload)
                if not notice_tpl_bytes:
                    st.error("貼り紙テンプレートの読み込みに失敗しました。パスまたはアップロードファイルを確認してください。")
            if generate_fax:
                fax_tpl_bytes = load_template_bytes(DEFAULT_FAX_TEMPLATE_PATH, fax_upload)
                if not fax_tpl_bytes:
                    st.error("FAXテンプレートの読み込みに失敗しました。パスまたはアップロードファイルを確認してください。")

            # テンプレートが1つも読めない場合は中断
            if (generate_notice and not notice_tpl_bytes) and (generate_fax and not fax_tpl_bytes):
                return

            files: List[tuple[str, bytes]] = []
            notice_count = 0
            fax_count = 0

            # 管理番号をキーに検索しやすく
            pm_idx = pm_view_df.set_index("管理番号")

            for ev in events:
                desc = safe_get(ev, "description") or ""
                summary = safe_get(ev, "summary") or ""

                mgmt = extract_assetnum(desc) or extract_assetnum(summary)
                if not mgmt:
                    continue

                mgmt_norm = normalize_str(mgmt)
                if mgmt_norm not in pm_idx.index:
                    continue

                pm_row = pm_idx.loc[mgmt_norm]

                start_dt = get_event_start_datetime(ev)
                end_dt = get_event_end_datetime(ev)

                # ---------- 貼り紙 ----------
                if generate_notice and notice_tpl_bytes:
                    kind = normalize_str(pm_row.get("貼り紙テンプレ種別", ""))
                    if kind == "自社":
                        reps = build_notice_replacements(
                            ev=ev,
                            pm_row=pm_row,
                            start_dt=start_dt,
                            end_dt=end_dt,
                            company_name=company_name,
                            company_tel=company_tel,
                        )
                        doc = Document(io.BytesIO(notice_tpl_bytes))
                        replace_placeholders_in_doc(doc, reps)
                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        fname = f"{mgmt_norm}_{display_value(pm_row.get('物件名', ''))}_貼り紙.docx"
                        files.append((fname, buf.getvalue()))
                        notice_count += 1

                # ---------- FAX ----------
                if generate_fax and fax_tpl_bytes:
                    fax1_flag = bool(normalize_str(pm_row.get("連絡方法_FAX1", "")))
                    fax2_flag = bool(normalize_str(pm_row.get("連絡方法_FAX2", "")))
                    if fax1_flag or fax2_flag:
                        reps = build_fax_replacements(
                            ev=ev,
                            pm_row=pm_row,
                            start_dt=start_dt,
                            end_dt=end_dt,
                            company_name=company_name,
                            company_tel=company_tel,
                            company_fax=company_fax,
                            contact_person=contact_person,
                        )
                        doc = Document(io.BytesIO(fax_tpl_bytes))
                        replace_placeholders_in_doc(doc, reps)
                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        fname = f"{mgmt_norm}_{display_value(pm_row.get('物件名', ''))}_FAX.docx"
                        files.append((fname, buf.getvalue()))
                        fax_count += 1

            if not files:
                st.info("条件に合致する貼り紙・FAXの作成対象がありませんでした。")
                return

            zip_bytes = build_zip_from_files(files)
            zip_name = f"notice_fax_{start_date.strftime('%Y%m%d')}_{end_date.strftime('%Y%m%d')}.zip"

            st.session_state["notice_fax_zip_bytes"] = zip_bytes
            st.session_state["notice_fax_zip_name"] = zip_name
            st.session_state["notice_fax_counts"] = {
                "notice": notice_count,
                "fax": fax_count,
            }

            st.success(
                f"貼り紙 {notice_count} 件、FAX {fax_count} 件の文書を作成しました。\n"
                "この下のボタンから ZIP ファイルとして一括ダウンロードできます。"
            )

    # 生成済みZIPのダウンロード
    zip_bytes = st.session_state.get("notice_fax_zip_bytes")
    if zip_bytes:
        zip_name = st.session_state.get("notice_fax_zip_name", "notice_fax.zip")
        counts = st.session_state.get("notice_fax_counts", {})
        label = "📦 貼り紙・FAX一括ダウンロード (ZIP)"
        if counts:
            label += f"  [貼り紙:{counts.get('notice', 0)} / FAX:{counts.get('fax', 0)}]"

        st.download_button(
            label=label,
            data=zip_bytes,
            file_name=zip_name,
            mime="application/zip",
            key="notice_fax_download_zip",
        )
