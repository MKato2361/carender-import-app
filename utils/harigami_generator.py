# utils/harigami_generator.py
from __future__ import annotations

import io
import os
import re
from datetime import datetime, timezone, timedelta
from typing import Dict, Optional, Tuple, Union

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 置換キー（テンプレ側のプレースホルダ → アプリ内キー）
PLACEHOLDERS: Dict[str, str] = {
    "［10月　19日（水）］": "DATE",
    "［10:00］": "START_TIME",
    "［11:00］": "END_TIME",
    "［物件名］": "NAME",
}

# 既定テンプレート対応表（/templates 配下）
# work_type → ファイル名
DEFAULT_TEMPLATE_MAP: Dict[str, str] = {
    "default": "harigami.docx",
    "点検": "harigami.docx",
    "検査": "kensa.docx",
    "有償工事": "paid.docx",
    "無償工事": "free.docx",
}

JST = timezone(timedelta(hours=9))

# --- Description からのタグ抽出（全角/半角対応） ---
_RE_WONUM = re.compile(r"[［\[]\s*作業指示書(?:番号)?\s*[：:]\s*([0-9A-Za-z\-]+)\s*[］\]]")
_RE_ASSET = re.compile(r"[［\[]\s*管理番号\s*[：:]\s*([0-9A-Za-z\-]+)\s*[］\]]")
_RE_NAME = re.compile(r"[［\[]\s*物件名\s*[：:]\s*(.+?)\s*[］\]]")


def extract_tags_from_description(desc: str) -> Dict[str, str]:
    tags: Dict[str, str] = {}
    if not desc:
        return tags
    m1 = _RE_WONUM.search(desc)
    if m1:
        tags["WONUM"] = m1.group(1).strip()
    m2 = _RE_ASSET.search(desc)
    if m2:
        tags["ASSETNUM"] = m2.group(1).strip()
    m3 = _RE_NAME.search(desc)
    if m3:
        tags["NAME"] = m3.group(1).strip()
    return tags


def _weekday_ja(dt: datetime) -> str:
    return ["月", "火", "水", "木", "金", "土", "日"][dt.weekday()]


def build_replacements_from_event(event: dict, summary: str, tags: Dict[str, str]) -> Dict[str, str]:
    """
    Google Calendar event からテンプレ置換データを生成
    """
    if "dateTime" in (event.get("start") or {}):
        start = event["start"]["dateTime"]
    elif "date" in (event.get("start") or {}):
        start = event["start"]["date"] + "T00:00:00+09:00"  # 終日扱い
    else:
        raise ValueError("invalid event start")

    if "dateTime" in (event.get("end") or {}):
        end = event["end"]["dateTime"]
    elif "date" in (event.get("end") or {}):
        end = event["end"]["date"] + "T23:59:59+09:00"
    else:
        raise ValueError("invalid event end")

    start_dt = _to_dt_jst(start)
    end_dt = _to_dt_jst(end)

    date_str = f"{start_dt.month}月{start_dt.day}日（{_weekday_ja(start_dt)}）"
    start_str = start_dt.strftime("%H:%M")
    end_str = end_dt.strftime("%H:%M")

    # NAME は Description の [物件名: ○○○] を最優先し、なければ引数 summary を使用
    name = (tags.get("NAME") or summary or "").strip() or "無題"

    replacements = {
        "DATE": date_str,
        "START_TIME": start_str,
        "END_TIME": end_str,
        "NAME": name,
    }

    # 使う/使わないはテンプレ依存だが、将来拡張用に同梱（未使用でも問題なし）
    if "WONUM" in tags:
        replacements["WONUM"] = tags["WONUM"]
    if "ASSETNUM" in tags:
        replacements["ASSETNUM"] = tags["ASSETNUM"]

    return replacements


def _to_dt_jst(val: str) -> datetime:
    # ISO8601 文字列を JST へ
    dt = datetime.fromisoformat(val.replace("Z", "+00:00"))
    return dt.astimezone(JST)


def _replace_text_across_runs(paragraph, search_text: str, replace_text: str):
    """
    段落内のテキストが複数 run に分割されていても search_text を探して置換する。
    """
    if not search_text:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if search_text not in full_text:
        return

    # まとめて置換
    new_full_text = full_text.replace(search_text, replace_text)

    # 既存の run をクリアして 1run に再構成
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_full_text
    else:
        run = paragraph.add_run(new_full_text)


def _replace_placeholders_preserve_format(paragraph, replacements: Dict[str, str]):
    """
    プレースホルダ単位で run のフォーマット（フォントサイズなど）を保存しつつ置換する。
    """
    full_text = paragraph.text
    if not full_text:
        return

    # まずは run を跨いでいる場合に備えた置換
    for ph, key in PLACEHOLDERS.items():
        if ph in full_text:
            _replace_text_across_runs(paragraph, ph, replacements.get(key, ""))

    # その後、run ごとにフォーマットを維持しながら微調整
    for ph, key in PLACEHOLDERS.items():
        if ph in paragraph.text:
            # run を走査してプレースホルダを含むものを探す
            for run in paragraph.runs:
                if ph in run.text:
                    original_font_size = run.font.size
                    original_bold = run.font.bold
                    original_italic = run.font.italic
                    original_underline = run.font.underline
                    original_color = run.font.color

                    run.text = run.text.replace(ph, replacements.get(key, ""))

                    if original_font_size:
                        run.font.size = original_font_size
                    if original_bold is not None:
                        run.font.bold = original_bold
                    if original_italic is not None:
                        run.font.italic = original_italic
                    if original_underline is not None:
                        run.font.underline = original_underline
                    if original_color is not None:
                        run.font.color.rgb = original_color.rgb


def _replace_placeholders_in_tables(doc, replacements: Dict[str, str]):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholders_preserve_format(paragraph, replacements)


def _replace_placeholders_comprehensive(doc, replacements: Dict[str, str]):
    for para in doc.paragraphs:
        if para.text.strip():
            _replace_placeholders_preserve_format(para, replacements)

    _replace_placeholders_in_tables(doc, replacements)

    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                _replace_placeholders_preserve_format(para, replacements)
        if section.footer:
            for para in section.footer.paragraphs:
                _replace_placeholders_preserve_format(para, replacements)


def generate_docx_from_template_like(
    template_like: Union[str, io.BytesIO],
    replacements: Dict[str, str],
    safe_title: str,
) -> Tuple[str, bytes]:
    """
    template_like: ファイルパス or BytesIO（アップロード）
    safe_title: 出力ファイル名のベース
    戻り値: (filename, content_bytes)
    """
    if isinstance(template_like, (io.BytesIO, io.BufferedReader)):
        template_like.seek(0)
        doc = Document(template_like)
    else:
        # パス
        doc = Document(template_like)

    _replace_placeholders_comprehensive(doc, replacements)

    # ファイル名安全化
    base = re.sub(r"[^\w\.\-]", "_", safe_title)
    base = re.sub(r"_{2,}", "_", base).strip("_") or "untitled_document"
    out_name = f"{base}.docx"

    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return out_name, mem.getvalue()
